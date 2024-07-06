import pandas as pd
import json
import requests
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Font, Border, Side

# Часть а
# Создаем три excel файла
file1 = pd.DataFrame({'A': [1, 2, 3], 'B': [4, 5, 6]})
file2 = pd.DataFrame({'A': [7, 8, 9], 'B': [10, 11, 12]})
file3 = pd.DataFrame({'A': [13, 14, 15], 'B': [16, 17, 18]})

file1.to_excel('1111.xlsx', index=False)
file2.to_excel('2222.xlsx', index=False)
file3.to_excel('3333.xlsx', index=False)

# Читаем данные из файлов
df1 = pd.read_excel('1111.xlsx')
df2 = pd.read_excel('2222.xlsx')
df3 = pd.read_excel('3333.xlsx')

# Объединяем данные в одну матрицу
combined_df = pd.concat([df1, df2, df3])

# Сортируем в порядке убывания
sorted_df = combined_df.sort_values(by='A', ascending=False)

# Записываем в новый excel файл с изменением шрифта и обертыванием в границы
with pd.ExcelWriter('sorted_combined.xlsx', engine='openpyxl') as writer:
    sorted_df.to_excel(writer, index=False)
    workbook = writer.book
    worksheet = writer.sheets['Sheet1']
    thin_border = Border(left=Side(style='thin'), 
                         right=Side(style='thin'), 
                         top=Side(style='thin'), 
                         bottom=Side(style='thin'))
    for row in worksheet.iter_rows():
        for cell in row:
            cell.font = Font(name='Arial', size=12)
            cell.border = thin_border

# Часть б
# Скачиваем данные с сайта и сохраняем в JSON файл
response = requests.get('https://jsonplaceholder.typicode.com/todos')
todos = response.json()

with open('todos.json', 'w') as f:
    json.dump(todos, f)

# Читаем JSON файл в массив словарей
with open('todos.json', 'r') as f:
    todos_data = json.load(f)

# Записываем каждый словарь в отдельный JSON файл
for i, todo in enumerate(todos_data):
    with open(f'todo_{i+1}.json', 'w') as f:
        json.dump(todo, f)

# Часть в
# Создаем Word файл с текстом "Hello Python"
doc = Document()
doc.add_paragraph('Hello Python')
doc.save('hello_python.docx')

# Читаем текст из файла, выделяем жирным и выводим на экран
doc = Document('hello_python.docx')
for para in doc.paragraphs:
    run = para.runs[0]
    run.bold = True
    print(para.text)

# Создаем абзац с текстом, изменяем шрифт и размер, записываем в новый Word файл
new_doc = Document()
new_paragraph = new_doc.add_paragraph()
run = new_paragraph.add_run('This is a new paragraph with custom font size.')
run.font.name = 'Calibri'
run.font.size = Pt(14)
new_doc.save('custom_paragraph.docx')